import io
import os
import uuid
import pickle
from datetime import datetime
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import re
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import sqlite3
import streamlit as st
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class FileParser:
    """Enhanced file parser with better error handling"""
    
    def __init__(self):
        os.makedirs('uploads/resumes', exist_ok=True)
        os.makedirs('uploads/job_descriptions', exist_ok=True)
    
    def save_uploaded_file(self, file_bytes, filename, file_type="resume"):
        """Save uploaded file and return path"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            file_extension = filename.split('.')[-1].lower()
            new_filename = f"{timestamp}_{unique_id}.{file_extension}"
            
            upload_dir = "uploads/resumes" if file_type == "resume" else "uploads/job_descriptions"
            file_path = os.path.join(upload_dir, new_filename)
            
            with open(file_path, 'wb') as f:
                f.write(file_bytes)
            
            logger.info(f"File saved successfully: {file_path}")
            return file_path
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            return None
    
    def extract_text_from_pdf(self, file_bytes):
        """Extract text from PDF files with enhanced parsing"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    # Clean up text
                    page_text = re.sub(r'\s+', ' ', page_text)
                    text += page_text + "\n"
            return text.strip() if text.strip() else ""
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_bytes):
        """Extract text from DOCX files"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip() if text.strip() else ""
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_image(self, file_bytes):
        """Extract text from images using OCR with enhanced settings"""
        try:
            image = Image.open(io.BytesIO(file_bytes))
            
            # Enhance image for better OCR
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Configure OCR for better accuracy
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            return text.strip() if text.strip() else ""
        except Exception as e:
            logger.error(f"Error reading image: {e}")
            return ""
    
    def parse_file(self, file_bytes, filename):
        """Parse any supported file type and extract text with enhanced error handling"""
        file_extension = filename.split('.')[-1].lower()
        
        try:
            if file_extension == 'pdf':
                text = self.extract_text_from_pdf(file_bytes)
            elif file_extension == 'docx':
                text = self.extract_text_from_docx(file_bytes)
            elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                text = self.extract_text_from_image(file_bytes)
            elif file_extension == 'txt':
                text = file_bytes.decode('utf-8', errors='ignore')
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
            
            # Validate extracted text
            if not text or len(text.strip()) < 10:
                logger.warning("No substantial text could be extracted from the file")
                return ""
                
            # Clean and normalize text
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
                
            return text
            
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {e}")
            return ""

class ResumeValidator:
    """Enhanced resume validator with better section detection"""
    
    def __init__(self):
        self.required_sections = ['education', 'experience', 'skills']
        self.optional_sections = ['summary', 'objective', 'projects', 'certifications', 'work']
    
    def validate_resume(self, text):
        """Check if resume has required sections and sufficient content"""
        if not text or len(text.strip()) < 50:
            return False, ["Resume text is too short or empty"]
            
        text_lower = text.lower()
        missing_sections = []
        
        for section in self.required_sections:
            if section not in text_lower:
                missing_sections.append(section)
        
        return len(missing_sections) == 0, missing_sections

class FeatureExtractor:
    """Enhanced feature extractor with better error handling"""
    
    def __init__(self, model=None):
        try:
            if model is None:
                self.transformer_model = SentenceTransformer('all-MiniLM-L6-v2')
            else:
                self.transformer_model = model
            logger.info("Transformer model loaded successfully")
        except Exception as e:
            logger.error(f"Error loading transformer model: {e}")
            self.transformer_model = None

    def extract_features(self, resume_text, jd_text):
        """Extract features for ML model prediction with enhanced error handling"""
        try:
            if not self.transformer_model:
                logger.error("Transformer model not loaded")
                return None, None, None
            
            if not resume_text or not jd_text:
                logger.error("Empty resume or JD text")
                return None, None, None
            
            # Get embeddings
            resume_embedding = self.transformer_model.encode([resume_text])[0]
            jd_embedding = self.transformer_model.encode([jd_text])[0]
            
            # Calculate cosine similarity
            cos_similarity = cosine_similarity([resume_embedding], [jd_embedding])[0][0] * 100
            
            # Calculate keyword match
            keyword_match = self._calculate_keyword_match(resume_text, jd_text)
            
            # Calculate skill overlap
            skill_overlap = self._calculate_skill_overlap(resume_text, jd_text)
            
            # Other features
            resume_length = len(resume_text.split())
            sections_count = self._count_sections(resume_text)
            
            features = {
                'keyword_match': keyword_match,
                'skill_overlap': skill_overlap,
                'resume_length': resume_length,
                'sections_count': sections_count,
                'cosine_similarity': cos_similarity
            }
            
            logger.info("Features extracted successfully")
            return features, resume_embedding, jd_embedding
            
        except Exception as e:
            logger.error(f"Error extracting features: {e}")
            return None, None, None

    def _calculate_keyword_match(self, resume_text, jd_text):
        """Calculate keyword match percentage between resume and JD"""
        try:
            vectorizer = CountVectorizer(stop_words='english', max_features=1000)
            vectors = vectorizer.fit_transform([resume_text, jd_text]).toarray()
            
            if np.sum(vectors[0]) == 0 or np.sum(vectors[1]) == 0:
                return 0
                
            cosine_sim = cosine_similarity(vectors[0].reshape(1, -1), vectors[1].reshape(1, -1))
            return cosine_sim[0][0] * 100
        except Exception as e:
            logger.warning(f"Error in keyword match calculation: {e}")
            return 0
    
    def _calculate_skill_overlap(self, resume_text, jd_text):
        """Calculate skill overlap percentage with enhanced skill list"""
        try:
            common_skills = [
                'python', 'java', 'javascript', 'typescript', 'sql', 'nosql', 'mongodb', 
                'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
                'machine learning', 'ai', 'data analysis', 'data science', 'tableau', 'power bi',
                'html', 'css', 'sass', 'bootstrap', 'tailwind', 'rest api', 'graphql',
                'linux', 'unix', 'windows', 'macos', 'agile', 'scrum', 'devops',
                'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
                'project management', 'time management', 'adaptability', 'creativity'
            ]
            
            resume_lower = resume_text.lower()
            jd_lower = jd_text.lower()
            
            resume_skills = [skill for skill in common_skills if skill in resume_lower]
            jd_skills = [skill for skill in common_skills if skill in jd_lower]
            
            if not jd_skills:
                return 0
            
            overlap = len(set(resume_skills).intersection(set(jd_skills)))
            return (overlap / len(jd_skills)) * 100
        except Exception as e:
            logger.warning(f"Error in skill overlap calculation: {e}")
            return 0
    
    def _count_sections(self, text):
        """Count number of important sections in resume with enhanced detection"""
        sections = ['education', 'experience', 'skills', 'work', 'projects', 'certifications', 'summary', 'objective']
        text_lower = text.lower()
        return sum(1 for section in sections if section in text_lower)

class DatabaseManager:
    """Enhanced database manager with better error handling"""
    
    def __init__(self, db_path='database/ats_db.sqlite'):
        self.db_path = db_path
        os.makedirs('database', exist_ok=True)
        self._init_db()
    
    def _init_db(self):
        """Initialize database tables with enhanced schema"""
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            
            # Create job_seekers table
            c.execute('''
                CREATE TABLE IF NOT EXISTS job_seekers (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    email TEXT UNIQUE NOT NULL,
                    resume_text TEXT,
                    resume_file_path TEXT,
                    skills TEXT,
                    experience TEXT,
                    education TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Create job_descriptions table
            c.execute('''
                CREATE TABLE IF NOT EXISTS job_descriptions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    company TEXT NOT NULL,
                    description_text TEXT,
                    jd_file_path TEXT,
                    required_skills TEXT,
                    experience_required TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Create matches table for storing ATS results
            c.execute('''
                CREATE TABLE IF NOT EXISTS matches (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    job_seeker_id INTEGER,
                    jd_id INTEGER,
                    ats_score REAL,
                    similarity_score REAL,
                    keyword_match REAL,
                    skill_overlap REAL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (job_seeker_id) REFERENCES job_seekers (id),
                    FOREIGN KEY (jd_id) REFERENCES job_descriptions (id)
                )
            ''')
            
            conn.commit()
            conn.close()
            logger.info("Database initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing database: {e}")
    
    def store_resume(self, name, email, resume_text, resume_file_bytes, resume_filename):
        """Store resume in database with enhanced error handling"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        try:
            file_parser = FileParser()
            resume_file_path = file_parser.save_uploaded_file(
                resume_file_bytes, resume_filename, "resume"
            )
            
            if not resume_file_path:
                return False
            
            # Extract skills and experience
            skills = self._extract_skills(resume_text)
            experience = self._extract_experience(resume_text)
            education = self._extract_education(resume_text)
            
            c.execute('''INSERT OR REPLACE INTO job_seekers 
                         (name, email, resume_text, resume_file_path, skills, experience, education)
                         VALUES (?, ?, ?, ?, ?, ?, ?)''',
                      (name, email, resume_text, resume_file_path, skills, experience, education))
            
            conn.commit()
            logger.info(f"Resume stored successfully for {email}")
            return True
        except sqlite3.IntegrityError:
            logger.warning(f"Resume with email {email} already exists")
            return False
        except Exception as e:
            logger.error(f"Error saving resume: {e}")
            return False
        finally:
            conn.close()
    
    def store_job_description(self, title, company, description_text, jd_file_bytes, jd_filename):
        """Store job description in database with enhanced error handling"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        try:
            file_parser = FileParser()
            jd_file_path = file_parser.save_uploaded_file(
                jd_file_bytes, jd_filename, "job_description"
            )
            
            if not jd_file_path:
                return False
            
            required_skills = self._extract_skills(description_text)
            experience_required = self._extract_experience(description_text)
            
            c.execute('''INSERT INTO job_descriptions 
                         (title, company, description_text, jd_file_path, required_skills, experience_required)
                         VALUES (?, ?, ?, ?, ?, ?)''',
                      (title, company, description_text, jd_file_path, required_skills, experience_required))
            
            conn.commit()
            logger.info(f"Job description stored successfully for {title} at {company}")
            return True
        except Exception as e:
            logger.error(f"Error saving job description: {e}")
            return False
        finally:
            conn.close()
    
    def get_all_resumes(self):
        """Get all resumes from database with error handling"""
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute('SELECT * FROM job_seekers ORDER BY created_at DESC')
            resumes = c.fetchall()
            conn.close()
            return resumes
        except Exception as e:
            logger.error(f"Error fetching resumes: {e}")
            return []
    
    def get_all_jds(self):
        """Get all job descriptions from database with error handling"""
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute('SELECT * FROM job_descriptions ORDER BY created_at DESC')
            jds = c.fetchall()
            conn.close()
            return jds
        except Exception as e:
            logger.error(f"Error fetching job descriptions: {e}")
            return []
    
    def store_match_result(self, job_seeker_id, jd_id, ats_score, similarity_score, keyword_match, skill_overlap):
        """Store ATS match results in database"""
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute('''INSERT INTO matches 
                         (job_seeker_id, jd_id, ats_score, similarity_score, keyword_match, skill_overlap)
                         VALUES (?, ?, ?, ?, ?, ?)''',
                      (job_seeker_id, jd_id, ats_score, similarity_score, keyword_match, skill_overlap))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error storing match result: {e}")
            return False
    
    def _extract_skills(self, text):
        """Enhanced skill extraction from text"""
        skills_keywords = [
            'Python', 'Java', 'SQL', 'JavaScript', 'Machine Learning', 'Data Analysis',
            'AWS', 'Docker', 'Communication', 'Teamwork', 'Problem Solving', 'Leadership',
            'Project Management', 'Agile', 'Scrum', 'Excel', 'PowerPoint', 'Word',
            'HTML', 'CSS', 'React', 'Angular', 'Vue', 'Node.js', 'Express', 'Django',
            'Flask', 'FastAPI', 'MongoDB', 'MySQL', 'PostgreSQL', 'Oracle', 'Git',
            'GitHub', 'Jenkins', 'Kubernetes', 'Linux', 'Windows', 'macOS'
        ]
        found_skills = [skill for skill in skills_keywords if skill.lower() in text.lower()]
        return ", ".join(found_skills[:10])
    
    def _extract_experience(self, text):
        """Enhanced experience extraction from text"""
        matches = re.findall(r'(\d+)\s*(?:years?|yrs?)', text.lower())
        if matches:
            return str(max([int(m) for m in matches]))
        else:
            # Enhanced inference from context
            text_lower = text.lower()
            if any(word in text_lower for word in ['senior', 'lead', 'manager', 'director', 'principal']):
                return "5"
            elif any(word in text_lower for word in ['mid-level', 'intermediate', 'experienced', 'ii', 'iii']):
                return "3"
            elif any(word in text_lower for word in ['junior', 'entry', 'associate', 'i']):
                return "1"
            else:
                return "2"  # Default assumption
    
    def _extract_education(self, text):
        """Enhanced education level extraction"""
        text_lower = text.lower()
        if 'phd' in text_lower or 'ph.d' in text_lower or 'doctorate' in text_lower:
            return "PhD"
        elif 'master' in text_lower or 'ms' in text_lower or 'm.s' in text_lower or 'mba' in text_lower:
            return "Master's Degree"
        elif 'bachelor' in text_lower or 'bs' in text_lower or 'b.s' in text_lower or 'ba' in text_lower or 'b.a' in text_lower:
            return "Bachelor's Degree"
        elif 'associate' in text_lower or 'diploma' in text_lower:
            return "Associate Degree/Diploma"
        else:
            return "Not Specified"